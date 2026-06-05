"""CV parsing — deterministic, offline text extraction + basic profile structuring.

Pipeline (Constitution Principles II, III, IV):
  1. validate size (<= MAX_CV_SIZE_BYTES) and format (pdf/docx) — reject otherwise;
  2. extract text deterministically (PyMuPDF for PDF, python-docx for DOCX) — reject
     documents with no extractable text (e.g. scanned/image-only PDFs; OCR is out of scope);
  3. build a basic structured profile (language es/en, emails, phones, links, skills).

Text extraction and structuring here are fully deterministic and require no network or LLM,
so tests run offline and results are reproducible. Any future LLM-based enrichment must go
through the provider abstraction and must never produce the numeric score.
"""
from __future__ import annotations

import io
import re
from dataclasses import asdict, dataclass, field

from app.core.config import settings

# --- errors ---------------------------------------------------------------------------


class CVParseError(Exception):
    """Base class for CV parsing failures."""


class UnsupportedFormatError(CVParseError):
    """File is neither a PDF nor a DOCX."""


class FileTooLargeError(CVParseError):
    """File exceeds the configured maximum size."""


class NoTextExtractedError(CVParseError):
    """No extractable text was found (e.g. an image-only/scanned PDF)."""


# --- structured result ----------------------------------------------------------------


@dataclass(slots=True)
class ParsedCV:
    language: str  # "es" | "en" | "unknown"
    emails: list[str] = field(default_factory=list)
    phones: list[str] = field(default_factory=list)
    links: list[str] = field(default_factory=list)
    skills: list[str] = field(default_factory=list)
    char_count: int = 0

    def to_dict(self) -> dict:
        return asdict(self)


# --- detection helpers ----------------------------------------------------------------

_PDF_EXT = "pdf"
_DOCX_EXT = "docx"
_DOCX_MIME = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/docx",
}


def detect_content_type(filename: str | None, declared: str | None) -> str | None:
    """Return 'pdf', 'docx', or None based on extension (preferred) then MIME type."""
    if filename and "." in filename:
        ext = filename.rsplit(".", 1)[-1].lower()
        if ext == _PDF_EXT:
            return "pdf"
        if ext == _DOCX_EXT:
            return "docx"
    declared = (declared or "").lower()
    if declared == "application/pdf":
        return "pdf"
    if declared in _DOCX_MIME:
        return "docx"
    return None


# --- text extraction ------------------------------------------------------------------


def _extract_pdf(data: bytes) -> str:
    import pymupdf  # lazy import; offline, no network

    try:
        doc = pymupdf.open(stream=data, filetype="pdf")
    except Exception as exc:  # corrupt / unreadable file
        raise UnsupportedFormatError("Could not read the PDF file") from exc
    try:
        return "\n".join(doc.load_page(i).get_text("text") for i in range(doc.page_count))
    finally:
        doc.close()


def _extract_docx(data: bytes) -> str:
    from docx import Document  # lazy import; offline, no network

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:  # corrupt / unreadable file
        raise UnsupportedFormatError("Could not read the DOCX file") from exc
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# --- language detection (es / en) -----------------------------------------------------

_ES_STOPWORDS = {
    "de", "la", "el", "en", "y", "los", "del", "las", "que", "con", "para", "una", "por",
    "experiencia", "años", "educación", "habilidades", "trabajo", "empresa", "actualidad",
}
_EN_STOPWORDS = {
    "the", "and", "of", "to", "in", "for", "with", "a", "an", "experience", "years",
    "education", "skills", "work", "company", "present", "summary",
}
_WORD_RE = re.compile(r"[a-záéíóúñü]+", re.IGNORECASE)


def detect_language(text: str) -> str:
    tokens = [t.lower() for t in _WORD_RE.findall(text)]
    es = sum(1 for t in tokens if t in _ES_STOPWORDS)
    en = sum(1 for t in tokens if t in _EN_STOPWORDS)
    if es == 0 and en == 0:
        return "unknown"
    return "es" if es >= en else "en"


# --- lightweight field extraction -----------------------------------------------------

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(r"\+?\d[\d\s().-]{7,}\d")
_LINK_RE = re.compile(r"https?://[^\s)]+", re.IGNORECASE)

# Known skill vocabulary (substring match, case-insensitive). Deterministic and offline.
_SKILL_VOCAB = [
    "python", "java", "javascript", "typescript", "fastapi", "django", "flask",
    "docker", "kubernetes", "postgresql", "mysql", "mongodb", "redis", "aws", "gcp",
    "azure", "react", "vue", "angular", "node", "sql", "git", "pytest", "linux",
    "go", "rust", "c++", "html", "css", "tailwind", "celery", "graphql",
]


def _dedup(items: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for it in items:
        key = it.lower()
        if key not in seen:
            seen.add(key)
            out.append(it)
    return out


def _structure(text: str) -> ParsedCV:
    lowered = text.lower()
    skills = [s for s in _SKILL_VOCAB if s in lowered]
    return ParsedCV(
        language=detect_language(text),
        emails=_dedup(_EMAIL_RE.findall(text)),
        phones=_dedup([p.strip() for p in _PHONE_RE.findall(text)]),
        links=_dedup(_LINK_RE.findall(text)),
        skills=skills,
        char_count=len(text),
    )


# --- public API -----------------------------------------------------------------------


def parse_cv(
    *,
    filename: str | None,
    declared_content_type: str | None,
    data: bytes,
    max_size: int | None = None,
) -> tuple[str, ParsedCV, str]:
    """Validate, extract and structure a CV.

    Returns ``(content_type, parsed, raw_text)``. Raises FileTooLargeError,
    UnsupportedFormatError, or NoTextExtractedError on rejection.
    """
    limit = settings.max_cv_size_bytes if max_size is None else max_size
    if len(data) > limit:
        raise FileTooLargeError(f"File exceeds the {limit}-byte limit")

    content_type = detect_content_type(filename, declared_content_type)
    if content_type is None:
        raise UnsupportedFormatError("Only PDF and DOCX files are supported")

    text = _extract_pdf(data) if content_type == "pdf" else _extract_docx(data)
    if not text.strip():
        raise NoTextExtractedError(
            "No extractable text found (image-only/scanned files are not supported)"
        )

    return content_type, _structure(text), text
